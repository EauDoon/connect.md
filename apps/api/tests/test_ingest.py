from __future__ import annotations

import asyncio
import io
import json
import os
import time
import zipfile
from pathlib import Path
from types import SimpleNamespace

import anyio
import pytest
from fastapi import HTTPException

import app.ingest as ingest_module
import app.ingest_worker as ingest_worker_module
from app.config import Settings
from app.ingest import SUPPORTED_UPLOAD_TYPES, _validate_file_signature
from app.ingest_worker import (
    _cleanup_pending_result_artifacts,
    _cleanup_stale_jobs,
    _convert_job,
    _process_request,
    _recover_orphaned_jobs,
)
from app.markdown import (
    PUBLIC_MARKDOWN_VALIDATION_DETAIL,
    canonical_document_max_utf8_bytes,
    client_template,
    prepare_client_document,
    split_markdown,
)
from tests.fixtures.ingest.generate import generate_fixtures


def test_binary_conversion_fixtures_are_deterministic_and_bounded(tmp_path) -> None:
    first = tmp_path / "first"
    second = tmp_path / "second"
    generate_fixtures(first)
    generate_fixtures(second)

    expected = {
        "valid.pdf",
        "valid.docx",
        "malformed.pdf",
        "malformed.docx",
        "oversized.pdf",
        "oversized.docx",
    }
    assert {path.name for path in first.iterdir()} == expected
    assert {path.name for path in second.iterdir()} == expected
    assert {path.name: path.read_bytes() for path in first.iterdir()} == {
        path.name: path.read_bytes() for path in second.iterdir()
    }
    assert (first / "valid.pdf").read_bytes().startswith(b"%PDF-")
    assert (first / "valid.docx").read_bytes().startswith(b"PK\x03\x04")
    assert not (first / "malformed.pdf").read_bytes().startswith(b"%PDF-")
    assert (first / "oversized.pdf").stat().st_size < 64 * 1024
    assert (first / "oversized.docx").stat().st_size < 64 * 1024


def test_pending_result_cleanup_is_exact_and_best_effort(tmp_path) -> None:
    pending = tmp_path / ".job.result.json.pending-123"
    unrelated = tmp_path / ".job.input.pending-123"
    pending.write_text("partial", encoding="utf-8")
    unrelated.write_text("keep", encoding="utf-8")

    _cleanup_pending_result_artifacts(tmp_path)

    assert not pending.exists()
    assert unrelated.read_text(encoding="utf-8") == "keep"


def test_pending_result_cleanup_runs_on_current_job_finalization(tmp_path) -> None:
    job_id = "finalize-pending"
    request_path = tmp_path / f"{job_id}.request.json"
    request_path.write_text(
        json.dumps(
            {
                "suffix": ".txt",
                "timeout_seconds": 5,
                "max_extracted_bytes": 1024,
            }
        ),
        encoding="utf-8",
    )
    (tmp_path / f"{job_id}.input").write_bytes(b"%PDF-test")
    pending = tmp_path / f".{job_id}.result.json.pending-123"
    pending.write_text("partial", encoding="utf-8")

    _process_request(request_path, tmp_path / ".worker-ready")

    assert not pending.exists()
    assert not request_path.exists()
    assert not (tmp_path / f"{job_id}.request.processing").exists()


def test_converter_failure_traceback_is_worker_only_and_redacted(
    tmp_path, monkeypatch, capsys
) -> None:
    private_detail = "diagnostic detail stays out of the protocol response"
    input_path = tmp_path / "source.pdf"
    output_path = tmp_path / "result.json"
    input_path.write_bytes(b"%PDF-test")

    def fail_conversion(_contents, _suffix, *, failure_reporter):
        try:
            raise RuntimeError(private_detail)
        except RuntimeError as exc:
            failure_reporter("markitdown", exc)
        raise HTTPException(
            status_code=422,
            detail={"message": "binary conversion failed", "warnings": ["RuntimeError"]},
        )

    monkeypatch.delenv("ORT_DISABLE_TELEMETRY", raising=False)
    monkeypatch.setattr(ingest_worker_module, "_convert_binary", fail_conversion)
    if os.name == "posix":
        monkeypatch.setattr(ingest_worker_module.os, "setsid", lambda: None)

    _convert_job(str(input_path), ".pdf", str(output_path), 1024)

    captured = capsys.readouterr()
    result = json.loads(output_path.read_text(encoding="utf-8"))
    assert os.environ["ORT_DISABLE_TELEMETRY"] == "1"
    assert captured.out == ""
    assert "event=ingest_conversion_failed converter=markitdown" in captured.err
    assert "Traceback (most recent call last)" in captured.err
    assert "fail_conversion" in captured.err
    assert private_detail not in captured.err
    assert result == {
        "ok": False,
        "status_code": 422,
        "message": "binary conversion failed",
        "warnings": ["RuntimeError"],
    }
    assert private_detail not in json.dumps(result)


def test_binary_signature_checks_reject_invalid_pdf_and_docx() -> None:
    settings = Settings(api_key_pepper="test-only-pepper-is-long-enough")
    with pytest.raises(HTTPException, match="PDF"):
        _validate_file_signature(b"not-a-pdf", ".pdf", settings)
    with pytest.raises(HTTPException, match="ZIP"):
        _validate_file_signature(b"not-a-docx", ".docx", settings)


def test_docx_archive_limits_are_checked_before_conversion() -> None:
    archive = io.BytesIO()
    with zipfile.ZipFile(archive, "w") as document:
        document.writestr("word/document.xml", "<document />")
        document.writestr("word/extra.xml", "<extra />")
    settings = Settings(api_key_pepper="test-only-pepper-is-long-enough", max_docx_entries=1)
    with pytest.raises(HTTPException, match="too many entries"):
        _validate_file_signature(archive.getvalue(), ".docx", settings)


def _real_docx() -> bytes:
    from docx import Document

    output = io.BytesIO()
    document = Document()
    document.add_heading("Ada Lovelace", level=1)
    document.add_paragraph("Python systems engineer")
    document.save(output)
    return output.getvalue()


def _real_pdf() -> bytes:
    from pypdf import PdfWriter
    from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

    output = io.BytesIO()
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
    )
    stream = DecodedStreamObject()
    stream.set_data(b"BT /F1 14 Tf 72 720 Td (Ada Lovelace Python engineer) Tj ET")
    page[NameObject("/Contents")] = writer._add_object(stream)
    writer.write(output)
    return output.getvalue()


@pytest.mark.parametrize(
    ("filename", "content_type", "payload"),
    [
        (
            "resume.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            _real_docx,
        ),
        ("resume.pdf", "application/pdf", _real_pdf),
        ("resume-unknown.pdf", "application/octet-stream", _real_pdf),
    ],
)
async def test_binary_endpoint_fails_closed_without_isolated_worker(
    api_client, filename: str, content_type: str, payload
) -> None:
    app, client = api_client
    app.state.settings.ingest_jobs_path = None
    response = await client.post(
        "/v1/ingest",
        data={"target_schema": "connect.md/resume"},
        files={"file": (filename, payload(), content_type)},
    )
    assert response.status_code == 503, response.text
    body = response.json()
    assert "draft_markdown" not in body
    assert body["detail"] == {
        "message": "binary conversion requires a configured isolated worker",
        "warnings": [],
        "provenance": {
            "source_type": filename.rsplit(".", 1)[-1],
            "converter": "none",
        },
    }


@pytest.mark.parametrize(
    ("filename", "content_type", "payload"),
    [
        (
            "resume.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            _real_docx,
        ),
        ("resume.pdf", "application/pdf", _real_pdf),
    ],
)
async def test_isolated_binary_endpoint_round_trip(
    api_client, tmp_path, filename: str, content_type: str, payload
) -> None:
    app, client = api_client
    jobs = tmp_path / "isolated-ingest-jobs"
    jobs.mkdir()
    heartbeat = jobs / ".worker-ready"
    heartbeat.touch()
    app.state.settings.ingest_jobs_path = jobs
    app.state.settings.ingest_timeout_seconds = 45

    async def process_one_request() -> None:
        deadline = time.monotonic() + 10
        while time.monotonic() < deadline:
            requests = list(jobs.glob("*.request.json"))
            if requests:
                await anyio.to_thread.run_sync(_process_request, requests[0], heartbeat)
                return
            await anyio.sleep(0.01)
        raise AssertionError("the API did not publish an isolated conversion request")

    worker = asyncio.create_task(process_one_request())
    try:
        response = await client.post(
            "/v1/ingest",
            data={"target_schema": "connect.md/resume"},
            files={"file": (filename, payload(), content_type)},
        )
    finally:
        await worker

    assert response.status_code == 200, response.text
    body = response.json()
    assert body["target_schema"] == "connect.md/resume"
    assert body["published"] is False
    assert "Ada Lovelace" in body["draft_markdown"]
    assert body["provenance"]["source_type"] == filename.rsplit(".", 1)[-1]
    assert body["provenance"]["converter"] in {
        "markitdown-local",
        "unstructured-local",
    }
    residue = [path for path in jobs.iterdir() if path.name != ".worker-ready"]
    assert all(path.name.endswith(".result.json") for path in residue)
    stale = time.time() - 1
    for path in residue:
        os.utime(path, (stale, stale))
    _cleanup_stale_jobs(jobs, maximum_age_seconds=0)
    assert sorted(path.name for path in jobs.iterdir()) == [".worker-ready"]


async def test_configured_binary_endpoint_requires_fresh_worker_heartbeat(
    api_client, tmp_path
) -> None:
    app, client = api_client
    jobs = tmp_path / "configured-without-worker"
    app.state.settings.ingest_jobs_path = jobs

    response = await client.post(
        "/v1/ingest",
        data={"target_schema": "connect.md/resume"},
        files={"file": ("resume.pdf", _real_pdf(), "application/pdf")},
    )

    assert response.status_code == 503, response.text
    assert response.json()["detail"] == {
        "message": "isolated conversion worker is unavailable",
        "warnings": [],
        "provenance": {"source_type": "pdf", "converter": "none"},
    }
    assert jobs.is_dir()
    assert list(jobs.iterdir()) == []


def test_isolated_wait_deadline_cleans_every_published_job_file(tmp_path, monkeypatch) -> None:
    heartbeat = tmp_path / ".worker-ready"
    heartbeat.touch()
    settings = Settings(
        api_key_pepper="test-only-pepper-is-long-enough",
        ingest_jobs_path=tmp_path,
        ingest_timeout_seconds=5,
    )
    ticks = iter((100.0, 116.0))
    monkeypatch.setattr(ingest_module, "monotonic", lambda: next(ticks))

    with pytest.raises(HTTPException) as caught:
        ingest_module._convert_isolated(b"%PDF-test", ".pdf", settings)

    assert caught.value.status_code == 503
    assert caught.value.detail == {
        "message": "isolated conversion worker did not return before the deadline",
        "warnings": [],
        "provenance": {"source_type": "pdf", "converter": "none"},
    }
    assert sorted(path.name for path in tmp_path.iterdir()) == [".worker-ready"]


def test_isolated_worker_timeout_result_preserves_503_envelope(tmp_path, monkeypatch) -> None:
    heartbeat = tmp_path / ".worker-ready"
    heartbeat.touch()
    settings = Settings(
        api_key_pepper="test-only-pepper-is-long-enough",
        ingest_jobs_path=tmp_path,
        ingest_timeout_seconds=5,
    )
    job_id = "timeout-result"
    monkeypatch.setattr(ingest_module, "uuid4", lambda: job_id)
    (tmp_path / f"{job_id}.result.json").write_text(
        json.dumps(
            {
                "ok": False,
                "status_code": 503,
                "message": "binary conversion exceeded the hard time limit",
                "warnings": [],
            }
        ),
        encoding="utf-8",
    )

    with pytest.raises(HTTPException) as caught:
        ingest_module._convert_isolated(b"%PDF-test", ".pdf", settings)

    assert caught.value.status_code == 503
    assert caught.value.detail == {
        "message": "binary conversion exceeded the hard time limit",
        "warnings": [],
        "provenance": {"source_type": "pdf", "converter": "none"},
    }
    assert sorted(path.name for path in tmp_path.iterdir()) == [".worker-ready"]


@pytest.mark.parametrize("source", ["true", "123", "null", "- foo", "2026-08-03"])
def test_plain_text_scalar_values_render_as_valid_yaml_strings(source: str) -> None:
    draft = client_template("profile", source)
    prepare_client_document(
        "profile",
        draft,
        document_id="3e811ba3-8d22-4aaf-a49e-b5b3a03eef0f",
        owner_id="owner-test",
        version=1,
    )


def test_imported_identifier_truncation_never_leaves_trailing_hyphen() -> None:
    draft = client_template("profile", "a" * 62 + " b")
    assert "handle: " + "a" * 62 + "\n" in draft
    prepare_client_document(
        "profile",
        draft,
        document_id="3e811ba3-8d22-4aaf-a49e-b5b3a03eef0f",
        owner_id="owner-test",
        version=1,
    )


def test_structured_import_preserves_resume_sections_without_contract_heading_collisions() -> None:
    source = """---
source: uploaded
---
# Ada Lovelace
Computing pioneer

## Professional Summary
Designed analytical systems.

## Work Experience
### Analytical Engine Society
- Designed the first published algorithm.

## Education
- Private study in mathematics

## Technical Skills
- Python
- Systems Design, Technical Writing
"""
    draft = client_template("resume", source)

    assert "name: Ada Lovelace" in draft
    assert "title: Computing pioneer" in draft
    assert "## Summary\n\nDesigned analytical systems." in draft
    assert "## Experience\n\n### Analytical Engine Society" in draft
    assert "## Education\n\n- Private study in mathematics" in draft
    frontmatter, _ = split_markdown(draft)
    assert [reference["label"] for reference in frontmatter["skills"]] == [
        "Python",
        "Systems Design",
        "Technical Writing",
    ]
    assert draft.count("# Ada Lovelace") == 1
    prepare_client_document(
        "resume",
        draft,
        document_id="3e811ba3-8d22-4aaf-a49e-b5b3a03eef0f",
        owner_id="owner-test",
        version=1,
    )


def test_plain_extracted_section_labels_are_structured() -> None:
    draft = client_template(
        "resume",
        "Ada Lovelace\nMathematician\nSummary\nComputing pioneer.\n"
        "Work Experience\nAnalytical Engine Society\nEducation\nPrivate mathematics study\n"
        "Skills\nPython, Algorithms",
    )
    assert "## Summary\n\nComputing pioneer." in draft
    assert "## Experience\n\nAnalytical Engine Society" in draft
    assert "## Education\n\nPrivate mathematics study" in draft
    frontmatter, _ = split_markdown(draft)
    assert [reference["label"] for reference in frontmatter["skills"]] == [
        "Python",
        "Algorithms",
    ]


@pytest.mark.parametrize("kind", ["profile", "resume"])
def test_default_ingest_template_is_deterministic_v2_and_discloses_only_source_fields(
    kind: str,
) -> None:
    source = "Ada Lovelace\nSystems engineer\nSkills\nPython, Systems Design"
    draft = client_template(kind, source)
    assert draft == client_template(kind, source)
    frontmatter, _ = split_markdown(draft)
    assert frontmatter["schema_version"] == 2
    assert frontmatter["work_modes"] == []
    assert frontmatter["availability"] == {"status": "not_disclosed"}
    assert frontmatter["public_representation"] == {"status": "not_disclosed"}
    assert frontmatter["contact"] == {"disclosure": "none"}
    assert frontmatter["location"]["label"] == "Not disclosed"
    assert frontmatter["occupations"][0]["label"] == "Not disclosed"
    assert all(
        reference["scheme"] == "connectmd-user" and reference["id"].startswith("connectmd-user-")
        for reference in [
            *frontmatter["occupations"],
            frontmatter["location"],
            *frontmatter["skills"],
            frontmatter["seniority"],
        ]
    )
    assert "Singapore" not in draft
    assert "available_now" not in draft
    prepare_client_document(
        kind,
        draft,
        document_id="3e811ba3-8d22-4aaf-a49e-b5b3a03eef0f",
        owner_id="owner-test",
        version=1,
    )


async def test_ingest_defaults_to_v2_and_accepts_only_explicit_v1_compatibility(api_client) -> None:
    _, client = api_client
    source = b"Ada Lovelace\nSystems engineer\nSkills\nPython"
    default = await client.post(
        "/v1/ingest",
        files={"file": ("profile.txt", source, "text/plain")},
    )
    assert default.status_code == 200, default.text
    default_frontmatter, _ = split_markdown(default.json()["draft_markdown"])
    assert default.json()["target_schema"] == "connect.md/profile"
    assert default.json()["provenance"]["schema_version"] == "2"
    assert default_frontmatter["schema_version"] == 2
    assert default_frontmatter["work_modes"] == []

    legacy = await client.post(
        "/v1/ingest",
        data={"target_schema": "connect.md/resume/v1"},
        files={"file": ("resume.txt", source, "text/plain")},
    )
    assert legacy.status_code == 200, legacy.text
    legacy_frontmatter, _ = split_markdown(legacy.json()["draft_markdown"])
    assert legacy.json()["target_schema"] == "connect.md/resume"
    assert legacy.json()["provenance"]["schema_version"] == "1"
    assert legacy_frontmatter["schema_version"] == 1


async def test_ingest_capabilities_match_source_contract_and_runtime_limits(api_client) -> None:
    app, client = api_client
    app.state.settings.ingest_jobs_path = None
    capability_document = (await client.get("/v1/capabilities")).json()
    capabilities = capability_document["ingestion"]
    assert capability_document["canonical_markdown"] == {
        "profile_resume_max_utf8_bytes": canonical_document_max_utf8_bytes(),
        "measurement": "final_rendered_utf8_bytes_after_lf_canonicalization",
        "json_schema_max_length_is_not_byte_proof": True,
    }
    sources = {item["extensions"][0]: item["mime_types"] for item in capabilities["formats"]}
    assert sources == {
        suffix: sorted(media_types) for suffix, media_types in SUPPORTED_UPLOAD_TYPES.items()
    }
    assert capabilities["limits"] == {
        "max_upload_bytes": app.state.settings.max_upload_bytes,
        "max_extracted_bytes": app.state.settings.max_extracted_bytes,
        "canonical_document_max_utf8_bytes": canonical_document_max_utf8_bytes(),
        "conversion_timeout_seconds": app.state.settings.ingest_timeout_seconds,
        "max_docx_entries": app.state.settings.max_docx_entries,
        "max_docx_uncompressed_bytes": app.state.settings.max_docx_uncompressed_bytes,
    }
    assert capabilities["binary_conversion"] == {
        "requires_isolated_worker": True,
        "worker_configured": False,
        "worker_heartbeat_ready": False,
        "unconfigured_behavior": "fail_closed",
    }

    jobs = app.state.settings.storage_path.parent / "capability-ingest-jobs"
    jobs.mkdir()
    heartbeat = jobs / ".worker-ready"
    heartbeat.touch()
    app.state.settings.ingest_jobs_path = jobs
    configured = (await client.get("/v1/capabilities")).json()["ingestion"]
    assert configured["binary_conversion"] == {
        "requires_isolated_worker": True,
        "worker_configured": True,
        "worker_heartbeat_ready": True,
        "unconfigured_behavior": "fail_closed",
    }
    stale = time.time() - 10
    os.utime(heartbeat, (stale, stale))
    stale_worker = (await client.get("/v1/capabilities")).json()["ingestion"]
    assert stale_worker["binary_conversion"]["worker_configured"] is True
    assert stale_worker["binary_conversion"]["worker_heartbeat_ready"] is False

    llms = await client.get("/llms-full.txt")
    assert llms.status_code == 200
    assert "GET /v1/capabilities" in llms.text
    assert str(canonical_document_max_utf8_bytes()) in llms.text


async def test_ingest_over_limit_draft_is_structured_unpublished_failure(api_client) -> None:
    _, client = api_client
    source = b"Ada Lovelace\nEngineer\n" + b"x" * canonical_document_max_utf8_bytes()
    response = await client.post(
        "/v1/ingest",
        data={"target_schema": "connect.md/profile"},
        files={"file": ("oversized.txt", source, "text/plain")},
    )
    assert response.status_code == 422, response.text
    body = response.json()
    assert "draft_markdown" not in body
    assert body["detail"]["message"] == PUBLIC_MARKDOWN_VALIDATION_DETAIL
    assert "canonical Profile/Resume Markdown exceeds" not in response.text
    assert body["detail"]["provenance"]["source_type"] == "txt"


def test_worker_recovers_interrupted_jobs_and_removes_expired_protocol_files(tmp_path) -> None:
    job_id = "recoverable"
    processing = tmp_path / f"{job_id}.request.processing"
    processing.write_text(json.dumps({"suffix": ".pdf"}), encoding="utf-8")
    (tmp_path / f"{job_id}.input").write_bytes(b"%PDF-test")

    _recover_orphaned_jobs(tmp_path)

    assert not processing.exists()
    assert (tmp_path / f"{job_id}.request.json").exists()

    stale = tmp_path / "stale.result.json"
    stale.write_text("{}", encoding="utf-8")
    old = time.time() - 600
    os.utime(stale, (old, old))
    _cleanup_stale_jobs(tmp_path, maximum_age_seconds=300)
    assert not stale.exists()


def test_worker_hard_timeout_terminates_job_and_returns_retryable_error(
    tmp_path, monkeypatch
) -> None:
    job_id = "timeout"
    request_path = tmp_path / f"{job_id}.request.json"
    request_path.write_text(
        json.dumps(
            {
                "suffix": ".pdf",
                "timeout_seconds": 5,
                "max_extracted_bytes": 1024,
            }
        ),
        encoding="utf-8",
    )
    input_path = tmp_path / f"{job_id}.input"
    input_path.write_bytes(b"%PDF-test")
    heartbeat = tmp_path / ".worker-ready"
    heartbeat.touch()

    class FakeProcess:
        pid = 12345
        alive = True
        started = False

        def start(self) -> None:
            self.started = True

        def is_alive(self) -> bool:
            return self.alive

        def join(self, _timeout: float | None = None) -> None:
            return None

    process = FakeProcess()

    def process_factory(*, target, args):
        assert target is ingest_worker_module._convert_job
        assert args == (str(input_path), ".pdf", str(tmp_path / f"{job_id}.result.json"), 1024)
        return process

    context = SimpleNamespace(Process=process_factory)
    monkeypatch.setattr(
        ingest_worker_module.multiprocessing,
        "get_context",
        lambda method: context if method == "spawn" else None,
    )
    ticks = iter((100.0, 106.0))
    monkeypatch.setattr(ingest_worker_module.time, "monotonic", lambda: next(ticks))
    terminated: list[FakeProcess] = []

    def terminate(candidate) -> None:
        candidate.alive = False
        terminated.append(candidate)

    monkeypatch.setattr(ingest_worker_module, "_terminate_process_tree", terminate)

    _process_request(request_path, heartbeat)

    assert process.started is True
    assert terminated == [process]
    assert json.loads((tmp_path / f"{job_id}.result.json").read_text(encoding="utf-8")) == {
        "ok": False,
        "status_code": 503,
        "message": "binary conversion exceeded the hard time limit",
        "warnings": [],
    }
    assert not request_path.exists()
    assert not input_path.exists()
    assert not (tmp_path / f"{job_id}.request.processing").exists()


@pytest.mark.parametrize("failure_stage", ["context", "process"])
def test_worker_startup_failure_returns_503_and_processes_next_request(
    tmp_path, monkeypatch, failure_stage
) -> None:
    heartbeat = tmp_path / ".worker-ready"
    heartbeat.touch()

    def create_job(job_id: str) -> Path:
        request_path = tmp_path / f"{job_id}.request.json"
        request_path.write_text(
            json.dumps(
                {
                    "suffix": ".pdf",
                    "timeout_seconds": 5,
                    "max_extracted_bytes": 1024,
                }
            ),
            encoding="utf-8",
        )
        (tmp_path / f"{job_id}.input").write_bytes(b"%PDF-test")
        return request_path

    first_request = create_job("startup-failure")
    second_request: Path | None = None
    context_calls = 0
    process_start_calls = 0

    class FakeProcess:
        def __init__(self, output_path: str) -> None:
            self.output_path = Path(output_path)

        def start(self) -> None:
            nonlocal process_start_calls
            process_start_calls += 1
            if failure_stage == "process" and process_start_calls == 1:
                raise RuntimeError("synthetic process-start failure")
            self.output_path.write_text(
                json.dumps(
                    {
                        "ok": True,
                        "text": "converted",
                        "converter": "fake",
                        "warnings": [],
                    }
                ),
                encoding="utf-8",
            )

        def is_alive(self) -> bool:
            return False

        def join(self, _timeout: float | None = None) -> None:
            return None

    def process_factory(*, target, args):
        assert target is ingest_worker_module._convert_job
        return FakeProcess(args[2])

    context = SimpleNamespace(Process=process_factory)

    def get_context(method: str):
        nonlocal context_calls
        assert method == "spawn"
        context_calls += 1
        if failure_stage == "context" and context_calls == 1:
            raise RuntimeError("synthetic process-start failure")
        return context

    monkeypatch.setattr(ingest_worker_module.multiprocessing, "get_context", get_context)

    _process_request(first_request, heartbeat)

    assert json.loads((tmp_path / "startup-failure.result.json").read_text(encoding="utf-8")) == {
        "ok": False,
        "status_code": 503,
        "message": "isolated worker rejected an invalid conversion job",
        "warnings": ["RuntimeError"],
    }
    assert not first_request.exists()
    assert not (tmp_path / "startup-failure.input").exists()

    second_request = create_job("after-startup-failure")
    _process_request(second_request, heartbeat)

    assert json.loads(
        (tmp_path / "after-startup-failure.result.json").read_text(encoding="utf-8")
    ) == {
        "ok": True,
        "text": "converted",
        "converter": "fake",
        "warnings": [],
    }
    assert context_calls == 2
    assert not second_request.exists()
    assert not (tmp_path / "after-startup-failure.input").exists()


def test_worker_supervision_failure_terminates_started_converter_child(
    tmp_path, monkeypatch
) -> None:
    job_id = "heartbeat-failure"
    request_path = tmp_path / f"{job_id}.request.json"
    request_path.write_text(
        json.dumps(
            {
                "suffix": ".pdf",
                "timeout_seconds": 5,
                "max_extracted_bytes": 1024,
            }
        ),
        encoding="utf-8",
    )
    input_path = tmp_path / f"{job_id}.input"
    input_path.write_bytes(b"%PDF-test")

    class FailingHeartbeat:
        def touch(self) -> None:
            raise OSError("heartbeat storage unavailable")

    class FakeProcess:
        pid = 12345
        alive = True
        started = False

        def start(self) -> None:
            self.started = True

        def is_alive(self) -> bool:
            return self.alive

        def join(self, _timeout: float | None = None) -> None:
            return None

    process = FakeProcess()

    def process_factory(*, target, args):
        assert target is ingest_worker_module._convert_job
        assert args == (str(input_path), ".pdf", str(tmp_path / f"{job_id}.result.json"), 1024)
        return process

    context = SimpleNamespace(Process=process_factory)
    monkeypatch.setattr(
        ingest_worker_module.multiprocessing,
        "get_context",
        lambda method: context if method == "spawn" else None,
    )
    monkeypatch.setattr(ingest_worker_module.time, "monotonic", lambda: 100.0)
    terminated: list[FakeProcess] = []

    def terminate(candidate) -> None:
        candidate.alive = False
        terminated.append(candidate)

    monkeypatch.setattr(ingest_worker_module, "_terminate_process_tree", terminate)

    _process_request(request_path, FailingHeartbeat())  # type: ignore[arg-type]

    assert process.started is True
    assert terminated == [process]
    assert json.loads((tmp_path / f"{job_id}.result.json").read_text(encoding="utf-8")) == {
        "ok": False,
        "status_code": 503,
        "message": "isolated worker rejected an invalid conversion job",
        "warnings": ["OSError"],
    }
    assert not request_path.exists()
    assert not input_path.exists()
    assert not (tmp_path / f"{job_id}.request.processing").exists()


async def test_invalid_binary_failure_is_structured(api_client) -> None:
    _, client = api_client
    response = await client.post(
        "/v1/ingest",
        data={"target_schema": "connect.md/profile"},
        files={"file": ("resume.pdf", b"not-a-pdf", "application/pdf")},
    )
    assert response.status_code == 422
    detail = response.json()["detail"]
    assert "PDF" in detail["message"]
    assert detail["warnings"] == []
    assert detail["provenance"] == {"source_type": "pdf", "converter": "none"}


async def test_text_ingest_accepts_mime_charset_parameter(api_client) -> None:
    app, client = api_client
    app.state.settings.ingest_jobs_path = None
    response = await client.post(
        "/v1/ingest",
        data={"target_schema": "connect.md/profile"},
        files={"file": ("profile.txt", b"Ada Lovelace\nEngineer", "text/plain; charset=utf-8")},
    )
    assert response.status_code == 200, response.text
    body = response.json()
    assert "# Ada Lovelace" in body["draft_markdown"]
    assert body["published"] is False
    assert body["provenance"] == {
        "source_type": "txt",
        "converter": "direct",
        "schema_version": "2",
    }
